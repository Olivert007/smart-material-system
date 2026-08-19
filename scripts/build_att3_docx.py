# -*- coding: utf-8 -*-
"""生成《附件3.参赛成果提交表》初版（填充表单 + 附页成果样例）。

用法: /usr/bin/python3 scripts/build_att3_docx.py
输出: /workspace/2026-07/附件3.参赛成果提交表-初版.docx
个人信息（姓名/出生年月/联系电话/单位）留空，由参赛人填写。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

TPL = Path("/workspace/2026-07/附件3.参赛成果提交表.docx")
OUT = Path("/workspace/2026-07/附件3.参赛成果提交表-初版.docx")

APP_SCENE = (
    "作品面向生产运营中的物资管理场景，解决备品备件、库存、需求、资产及出入库流水等台账"
    "\u201c数据多源异构、格式口径不一、人工查数难\u201d的现实痛点。各单位台账以 Excel/CSV 分散存放，"
    "字段口径不一致、数据质量参差不齐，人工合并清洗耗时费力且不可追溯。"
    "搭建必要性：一是物资数据治理标准化是精细化管理的基础；二是传统人工治理成本高、易出错、无法审计；"
    "三是数据涉及内部业务信息，必须全程保密合规。"
    "系统投入使用后，可自动完成台账解析、规则清洗、出入库流水智能拆解、主数据映射等治理工作，"
    "支持自然语言查询库存、需求、资产等指标，过程可审计、结果可回滚。"
    "预期显著降低治理人力成本、提升数据质量与查数效率，实现物资数据\u201c一次治理、长期受益\u201d，"
    "并为同类多源数据治理提供可复制样板。"
)

DESIGN = (
    "核心需求：打通物资领域多源异构台账，实现\u201c上传—治理—发布—问答\u201d全链路智能化，且全程数据不出本机。"
    "功能设计：①数据接入——Excel/CSV 上传，自动解析多 Sheet 结构、留存原始证据；"
    "②AI 治理——本地大模型完成字段画像、映射建议（向量召回）、出入库流水智能拆解，规则清洗全量覆盖，"
    "人工确认门槛保障准确性；③可信发布——唯一写入入口、幂等发布、数据血缘与审计时间线、一键备份；"
    "④智能问答——自然语言问数（Text2SQL），指标模板优先、只读查询、图表展示；"
    "⑤运维监控——模型状态、任务队列、评测集持续评估。"
    "模型 API 接入：全部模型本地部署，vLLM 提供 OpenAI 兼容接口走内网回环——"
    "生成模型 Qwen3.6-27B（主生成）、Embedding 模型 Qwen3-Embedding-0.6B（映射召回），"
    "数据不出本机，完全符合保密红线；多模型按场景路由，评测集把关启用门槛。"
    "技术栈与工程：FastAPI + Vue3/Element Plus + SQLite（元数据/任务）+ DuckDB（分析库）+ "
    "Parquet（证据层）+ vLLM；Docker Compose 分服务部署。项目组成：app（API/服务/仓储/Worker）、"
    "frontend（前端）、deploy（容器编排）、scripts（运维与评测）、tests（自动化测试）。"
    "关键截图见本表下方附页与视频（3—4 张）。"
)


def set_cell(cell, text: str) -> None:
    cell.text = text


def tick(cell, option: str) -> None:
    t = cell.text
    t = t.replace("\u25a1 " + option, "\u25a0 " + option)  # □ → ■
    cell.text = t


def main() -> int:
    doc = Document(str(TPL))
    t = doc.tables[0]

    # R1 个人信息（留空）
    set_cell(t.rows[1].cells[1], "")
    set_cell(t.rows[1].cells[3], "")
    set_cell(t.rows[1].cells[7], "")

    # R2 单位（留空）；数字化工作=否
    set_cell(t.rows[2].cells[1], "")
    tick(t.rows[2].cells[7], "否")

    # R3 作品名称
    set_cell(t.rows[3].cells[1], "智能物资数据管理系统")

    # R4 应用场景 → 生产运营
    tick(t.rows[4].cells[1], "生产运营")

    # R5 数据来源 → 使用内部资料且已脱敏
    tick(t.rows[5].cells[1], "使用内部资料且已脱敏")

    # R6 使用平台 → Trae 平台；同意推广 → 是
    tick(t.rows[6].cells[1], "Trae 平台")
    tick(t.rows[6].cells[7], "是")

    # R7 应用场景说明
    set_cell(t.rows[7].cells[1], APP_SCENE)

    # R8 方案设计与关键配置（含截图占位）
    set_cell(t.rows[8].cells[1], DESIGN + "\n\n【关键截图（3—4 张）】\n1. 首页总览\n2. 治理中心（映射/流水/勾稽）\n3. 问数助手\n4. 运维面板/模型API")

    # ---------- 附页：成果样例 ----------
    doc.add_paragraph()
    h = doc.add_paragraph("附件3-附页 成果样例")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs:
        r.bold = True
        r.font.size = Pt(14)

    p = doc.add_paragraph(
        "需提供一组完整案例素材，包含1份输入内容、1份对应完整的输出结果。（随附件3上报，页数不限）"
    )
    p.paragraph_format.first_line_indent = Pt(21)

    doc.add_paragraph("一、输入内容：演示用物资台账（脱敏样例）.xlsx（虚构站点/人员/物料，保留真实结构）")
    for line in [
        "维护材料（21行）：序号、名称、品牌型号规格、现有库存(=初始+入库-出库)、单位、存放位置、保管人、初始库存、入库记录（时间/经手人/编码）、入库数量、出库记录（时间/领用人/用途）、出库数量、最低库存阈值",
        "备品备件（21行）：物资名称、现有库存量、入库/出库记录、所属系统、项目名称、消耗计划、物资来源、定额数量",
        "应急备汛物资（12行）：区域、物资编码、物资名称、定额数量、现有数量、存放货位、管理员",
        "公用工器具（15行）：物资名称、资产编码、数量、购买日期、工器具来源、是否仪器仪表",
        "脏数据特征：空字段、混合日期格式、出入库记录长文本、部分行\u201c现有库存≠初始+入库-出库\u201d",
    ]:
        doc.add_paragraph(line)

    doc.add_paragraph("二、输出结果")
    doc.add_paragraph("输出1：标准库存表 fact_inventory（发布49行，示例）：")
    inv_rows = [
        ["维护材料", "插线板/通用插线板 3米", "10", "12", "个", "云岭电站右岸中控室", "王强"],
        ["维护材料", "光模块/千兆光模块 SFP", "1", "3", "个", "望湖大厦6楼通信材料室", "孙浩"],
        ["维护材料", "网线/超五类网线", "11", "5", "米", "云岭电站右岸中控室", "孙浩"],
        ["维护材料", "熔纤盘/标准熔纤盘 12芯", "10", "7", "个", "青岩水电站东区仓库", "陈静"],
        ["备品备件", "整流模块/标准整流模块 48V/40A", "4", "2", "个", "锦澜电站左岸厂房仓库", "张伟"],
    ]
    _table(doc, ["类别", "物资名称/规格", "现有库存", "初始库存", "单位", "存放位置", "保管人"], inv_rows)

    doc.add_paragraph("输出2：出入库流水表 fact_stock_flow（发布48段，L1规则拆解；配置命中67、物料对齐22）：")
    flow_rows = [
        ["OUT", "2026-02-01", "2", "个", "调度大厅备用", "L1", "维护材料"],
        ["IN", "2026-04-03", "4", "个", "年度采购", "L1", "维护材料"],
        ["OUT", "2026-03-01", "3", "个", "调度大厅备用", "L1", "维护材料"],
        ["OUT", "2026-04-04", "2", "米", "现场检修使用", "L1", "维护材料"],
        ["IN", "2026-06-05", "6", "根", "年度采购", "L1", "维护材料"],
    ]
    _table(doc, ["流水类型", "发生日期", "数量", "单位", "用途", "拆解级别", "来源Sheet"], flow_rows)
    doc.add_paragraph("汇总：IN 23段/70数量；OUT 25段/71数量。")

    doc.add_paragraph("输出3：勾稽差异清单（自动暴露口径不一致39条，示例）：")
    rec_rows = [
        ["扎带", "2", "5", "-3", "-10", "-7", "mismatch"],
        ["潜水泵", "5", "0", "5", "0", "-5", "inv_only"],
    ]
    _table(doc, ["物资名称", "现有库存", "初始库存", "应有净变动", "流水净变动", "差异", "类型"], rec_rows)

    doc.add_paragraph("输出4：自然语言问数结果（Text2SQL）：")
    doc.add_paragraph("问：各存放位置库存数量排名 → 生成SQL：SELECT location, COUNT(*) AS n, SUM(stock_qty) AS total_qty FROM fact_inventory GROUP BY location ORDER BY total_qty DESC LIMIT 5")
    q_rows = [
        ["望湖大厦6楼通信材料室", "7", "47"],
        ["锦澜电站左岸厂房仓库", "9", "47"],
        ["青岩水电站东区仓库", "7", "39"],
        ["青岩水电站调度大楼", "6", "36"],
        ["望湖大厦负一楼材料室", "5", "34"],
    ]
    _table(doc, ["存放位置", "物资种类数", "库存总量"], q_rows)
    doc.add_paragraph("问：入库和出库分别多少条、合计多少数量 → SELECT flow_type, COUNT(*) AS n, SUM(quantity) AS qty FROM fact_stock_flow GROUP BY flow_type")
    q2_rows = [["IN", "23", "70"], ["OUT", "25", "71"]]
    _table(doc, ["流水类型", "条数", "数量合计"], q2_rows)

    doc.add_paragraph("复现方式：python3 scripts/build_demo_env.py 重建干净演示库后可在系统内复现上述结果。")

    doc.save(str(OUT))
    print("已生成:", OUT)
    return 0


def _table(doc, headers, rows):
    tb = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        tb.style = "Table Grid"
    except KeyError:
        pass
    for j, htext in enumerate(headers):
        c = tb.rows[0].cells[j]
        c.text = htext
        for r in c.paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            tb.rows[i].cells[j].text = str(v)


if __name__ == "__main__":
    raise SystemExit(main())
