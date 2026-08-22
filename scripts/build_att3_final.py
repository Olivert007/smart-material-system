# -*- coding: utf-8 -*-
"""生成《附件3.参赛成果提交表》完善版（正式模板填充 + 格式化 + 附页成果样例）。

用法: /usr/bin/python3 scripts/build_att3_final.py
前置: 先运行 python3 scripts/build_demo_env.py 重建干净演示库（demo_data/runtime）。
输出: /workspace/2026-07/附件3.参赛成果提交表-完善版.docx
个人信息（姓名/出生年月/联系电话/单位）留空，由参赛人填写。

数据口径（2026-08-19 基于新模板脱敏样例实测）：
- 输入：单 sheet 综合台账 909 行 / 19 列（维护材料/备品备件/低值易耗/个人工器具混合在一表）
- 输出：fact_inventory 662 行；fact_stock_flow 136 段（L1 纯规则拆解，流水文本段 244、配置命中 748、物料对齐 344、待人工确认 108）
- 本表不再包含「勾稽差异」内容（该特性已下线，口径与《成果样例（附页）.md》/演示视频保持一致）
"""
from __future__ import annotations

import json
import re
import sqlite3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

TPL = Path("/workspace/2026-07/附件3.参赛成果提交表.docx")
OUT = Path("/workspace/2026-07/附件3.参赛成果提交表-完善版.docx")
ROOT = Path(__file__).resolve().parents[1]
RUNTIME = ROOT / "demo_data" / "runtime"
DUCKDB = RUNTIME / "material.duckdb"
META = RUNTIME / "meta.sqlite"
SAMPLE_NAME = "desensitized-sample.xlsx"  # 本地 demo_data/samples/，不入库

APP_SCENE = (
    "作品面向生产运营中的物资管理场景，解决备品备件、库存、需求、资产及出入库流水等台账"
    "\u201c数据多源异构、格式口径不一、人工查数难\u201d的现实痛点。各单位台账以 Excel/CSV 分散存放，"
    "字段口径不一致、数据质量参差不齐，人工合并清洗耗时费力且不可追溯。"
    "搭建必要性：一是物资数据治理标准化是精细化管理的基础；二是传统人工治理成本高、易出错、无法审计；"
    "三是数据涉及内部业务信息，必须全程保密合规。"
    "系统投入使用后，可自动完成台账解析、规则清洗、出入库流水智能拆解、主数据映射等治理工作，"
    "支持自然语言查询库存、需求、资产等指标，过程可审计、结果可回滚。"
    "预期显著降低治理人力成本、提升数据质量与查数效率，实现物资数据\u201c一次治理、长期受益\u201d，"
    "并为同类多源数据治理提供可复制样板。"
)

DESIGN = (
    "核心需求：打通物资领域多源异构台账，实现\u201c上传—治理—发布—问答\u201d全链路智能化，且全程数据不出本机。"
    "功能设计：①数据接入——Excel/CSV 上传，自动解析多 Sheet/单 Sheet 混合结构、留存原始证据；"
    "②AI 治理——本地大模型完成字段画像、映射建议（向量召回）、出入库流水智能拆解，规则清洗全量覆盖，"
    "人工确认门槛保障准确性；③可信发布——唯一写入入口、幂等发布、数据血缘与审计时间线、一键备份；"
    "④智能问答——自然语言问数（Text2SQL），指标模板优先、只读查询、图表展示；"
    "⑤运维监控——模型状态、任务队列、评测集持续评估。"
    "模型 API 接入：全部模型本地部署，vLLM 提供 OpenAI 兼容接口走内网回环——"
    "生成模型 Qwen3.6-27B（主生成）、Embedding 模型 Qwen3-Embedding-0.6B（映射召回），"
    "数据不出本机，完全符合保密红线；多模型按场景路由，评测集把关启用门槛。"
    "技术栈与工程：FastAPI + Vue3/Element Plus + SQLite（元数据/任务）+ DuckDB（分析库）+ "
    "Parquet（证据层）+ vLLM；Docker Compose 分服务部署。项目组成：app（API/服务/仓储/Worker）、"
    "frontend（前端）、deploy（容器编排）、scripts（运维与评测）、tests（自动化测试）。"
    "关键截图见本表下方附页与视频（3—4 张）。"
)

SCREENSHOTS = (
    "【关键截图（3—4 张）】\n"
    "1. 首页总览：业务指标卡片与免责声明\n"
    "2. 治理中心：字段映射建议 / 出入库流水智能拆解（待确认队列）\n"
    "3. 问数助手：自然语言问数结果\n"
    "4. 运维面板：本地模型 API 接入与任务队列"
)


def _set_run_font(run, size=10.5, bold=False, font_zh="宋体"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), font_zh)


def set_cell(cell, text: str, size=10.5, bold=False, align=None) -> None:
    """写入单元格文本，统一字体字号；\\n 保留为单元格内换行。"""
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            _set_run_font(r, size=size, bold=bold)
        if align is not None:
            p.alignment = align
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def tick(cell, option: str) -> bool:
    """把 □ 选项勾选为 ■ 选项（兼容有无空格）。返回是否命中。"""
    t = cell.text
    for opt in (option, option.replace(" ", ""), option.replace(" ", " ")):
        pat = r"\u25a1\s*" + re.escape(opt)
        if re.search(pat, t):
            t = re.sub(pat, "\u25a0 " + option, t, count=1)
            cell.text = t
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_font(r)
            return True
    return False


def _table(doc, headers, rows, header_size=9, cell_size=9):
    tb = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        tb.style = "Table Grid"
    except KeyError:
        pass
    for j, htext in enumerate(headers):
        c = tb.rows[0].cells[j]
        set_cell(c, htext, size=header_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            set_cell(tb.rows[i].cells[j], str(v), size=cell_size)
    return tb


def _num(v) -> str:
    if v is None:
        return ""
    f = float(v)
    return str(int(f)) if f == int(f) else f"{f:g}"


def load_stats() -> dict:
    """从干净演示库读取权威统计（需先运行 build_demo_env.py）。"""
    import duckdb

    if not DUCKDB.exists() or not META.exists():
        raise RuntimeError("缺少 demo_data/runtime，请先运行: python3 scripts/build_demo_env.py")
    con = duckdb.connect(str(DUCKDB), read_only=True)
    s: dict = {}
    s["inv_rows"] = con.execute("SELECT COUNT(*) FROM fact_inventory").fetchone()[0]
    s["inv_qty"] = _num(con.execute("SELECT SUM(stock_qty) FROM fact_inventory").fetchone()[0])
    s["inv_by_cat"] = [
        (c, n, _num(q))
        for c, n, q in con.execute(
            "SELECT category, COUNT(*), SUM(stock_qty) FROM fact_inventory "
            "GROUP BY category ORDER BY 2 DESC"
        ).fetchall()
    ]
    s["inv_sample"] = [
        (c, name, _num(q), _num(quota), u, loc)
        for c, name, q, quota, u, loc in con.execute(
            "SELECT i.category, m.material_name, i.stock_qty, i.quota_qty, i.unit, i.location "
            "FROM fact_inventory i JOIN dim_material m ON m.material_id = i.material_id "
            "WHERE i.stock_qty > 0 "
            "QUALIFY ROW_NUMBER() OVER (PARTITION BY i.category ORDER BY i.stock_qty DESC) = 1 "
            "ORDER BY i.stock_qty DESC LIMIT 5"
        ).fetchall()
    ]
    s["flow_rows"] = con.execute("SELECT COUNT(*) FROM fact_stock_flow").fetchone()[0]
    s["flow_qty"] = _num(con.execute("SELECT SUM(quantity) FROM fact_stock_flow").fetchone()[0])
    s["flow_sample"] = [
        (ft, d, _num(q), u, p, pp, lv)
        for ft, d, q, u, p, pp, lv in con.execute(
            "SELECT flow_type, flow_date, quantity, unit, person, purpose, parse_level "
            "FROM fact_stock_flow "
            "WHERE person IS NOT NULL AND person != '' AND person != '.' "
            "AND length(purpose) > 2 AND purpose NOT LIKE '%.%' "
            "ORDER BY flow_date DESC LIMIT 5"
        ).fetchall()
    ]
    s["q_loc"] = [
        (loc, n, _num(q))
        for loc, n, q in con.execute(
            "SELECT location, COUNT(*), SUM(stock_qty) FROM fact_inventory "
            "GROUP BY location ORDER BY 3 DESC LIMIT 5"
        ).fetchall()
    ]
    s["flow_by_type"] = [
        (ft, n, _num(q))
        for ft, n, q in con.execute(
            "SELECT flow_type, COUNT(*), SUM(quantity) FROM fact_stock_flow GROUP BY 1"
        ).fetchall()
    ]
    con.close()

    mc = sqlite3.connect(str(META))
    mc.row_factory = sqlite3.Row
    s["flow_pending"] = mc.execute(
        "SELECT COUNT(*) FROM flow_pending WHERE status='pending'"
    ).fetchone()[0]
    s["write_audit"] = mc.execute("SELECT COUNT(*) FROM write_audit").fetchone()[0]
    dry = None
    for r in mc.execute(
        "SELECT dry_run_json FROM staging_record WHERE target_domain='stock_flow' "
        "ORDER BY created_at DESC LIMIT 1"
    ).fetchall():
        dry = json.loads(r["dry_run_json"] or "{}")
    mc.close()
    fp = (dry or {}).get("flow_parse") or {}
    s["segments"] = fp.get("segments", 0)
    s["config_hits"] = fp.get("config_hits", 0)
    s["material_aligned"] = fp.get("material_aligned", 0)
    s["l1"] = fp.get("L1", 0)
    return s


def main() -> int:
    s = load_stats()
    doc = Document(str(TPL))
    t = doc.tables[0]

    # R1 个人信息（留空，参赛人填写）
    set_cell(t.rows[1].cells[1], "")
    set_cell(t.rows[1].cells[3], "")
    set_cell(t.rows[1].cells[7], "")

    # R2 单位（留空）；数字化工作=否
    set_cell(t.rows[2].cells[1], "")
    ok = tick(t.rows[2].cells[7], "否")
    assert ok, "R2 数字化勾选失败"

    # R3 作品名称
    set_cell(t.rows[3].cells[1], "智能物资数据管理系统", bold=True)

    # R4 应用场景 → 生产运营
    ok = tick(t.rows[4].cells[1], "生产运营")
    assert ok, "R4 应用场景勾选失败"

    # R5 数据来源 → 使用内部资料且已脱敏
    ok = tick(t.rows[5].cells[1], "使用内部资料且已脱敏")
    assert ok, "R5 数据来源勾选失败"

    # R6 使用平台 → Trae平台；同意推广 → 是
    ok = tick(t.rows[6].cells[1], "Trae平台")
    assert ok, "R6 使用平台勾选失败"
    ok = tick(t.rows[6].cells[7], "是")
    assert ok, "R6 推广勾选失败"

    # R7 应用场景说明（约 300 字，符合 100-500 字要求）
    set_cell(t.rows[7].cells[1], APP_SCENE)

    # R8 方案设计与关键配置（约 700 字，符合 200-800 字要求）
    set_cell(t.rows[8].cells[1], DESIGN + "\n\n" + SCREENSHOTS)

    # ---------- 附页：成果样例（输入 + 输出） ----------
    doc.add_paragraph()
    h = doc.add_paragraph("附件3-附页 成果样例")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs:
        _set_run_font(r, size=14, bold=True, font_zh="黑体")

    p = doc.add_paragraph(
        "需提供一组完整案例素材，包含1份输入内容、1份对应完整的输出结果。（随附件3上报，页数不限）"
    )
    for r in p.runs:
        _set_run_font(r, size=10.5)

    h1 = doc.add_paragraph(f"一、输入内容（1 份）：{SAMPLE_NAME}")
    for r in h1.runs:
        _set_run_font(r, size=10.5, bold=True)
    for line in [
        "单 Sheet 综合台账（新模板），909 行 / 19 列；Sheet 名：演示区域ZW物资台账（新模板）。",
        "表头：序号、物资种类、物资名称、品牌型号规格、定额数量、临时储存数量（ZW）、公司仓库数量、"
        "现有库存、单位、存放位置、物资来源、物资储存时间（ZW）、消耗计划、出库记录（ZW）、出库数量、"
        "入库记录、入库数量、剩余临时储存数量（ZW）、备注。",
        "业务范围：维护材料 / 备品备件 / 低值易耗 / 个人工器具四类物资混合在一张综合表。",
        "脏数据特征：表头前的标题行与示例行（序号「例」）、出库记录为长文本（含时间/人员/数量/用途，"
        "如\u201c2025年10月，徐刚领用3个，用于机房音频配线\u201d）、入库记录列全空、"
        "单位写法不一（块/快）、品名重复、空字段与口径不一致。",
    ]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            _set_run_font(r, size=10.5)

    h2 = doc.add_paragraph("二、输出结果（对应上述输入的完整输出）")
    for r in h2.runs:
        _set_run_font(r, size=10.5, bold=True)

    p = doc.add_paragraph(
        f"输出1：标准库存表 fact_inventory（发布 {s['inv_rows']} 行，库存总量 {s['inv_qty']}；"
        "质量预检 0 阻断问题，示例 5 行）："
    )
    for r in p.runs:
        _set_run_font(r, size=10.5)
    inv_rows = [[c, name, q, quota, u, loc] for c, name, q, quota, u, loc in s["inv_sample"]]
    _table(doc, ["类别", "物资名称", "现有库存", "定额数量", "单位", "存放位置"], inv_rows)
    p = doc.add_paragraph(
        "按类别："
        + "；".join(f"{c} {n} 行 / {q}" for c, n, q in s["inv_by_cat"])
        + "。全部字段统一映射、带来源血缘，可直接被上层指标与问数消费。"
    )
    for r in p.runs:
        _set_run_font(r, size=10.5)

    p = doc.add_paragraph(
        f"输出2：出入库流水表 fact_stock_flow（发布 {s['flow_rows']} 段，合计 {s['flow_qty']}；"
        f"流水文本段 {s['segments']}、配置命中 {s['config_hits']}、物料对齐 {s['material_aligned']}、"
        f"待人工确认 {s['flow_pending']} 条，示例 5 段）："
    )
    for r in p.runs:
        _set_run_font(r, size=10.5)
    flow_rows = [
        [ft, d, q, u, p_, pp, lv] for ft, d, q, u, p_, pp, lv in s["flow_sample"]
    ]
    _table(doc, ["类型", "日期", "数量", "单位", "领用人", "用途", "拆解级别"], flow_rows)
    p = doc.add_paragraph(
        "汇总："
        + "；".join(f"{ft} {n} 段 / {q}" for ft, n, q in s["flow_by_type"])
        + "。规则路径 L1 自动拆解（日期/数量/人员/用途齐全），无法完全确认的文本进入待确认队列人工复核。"
    )
    for r in p.runs:
        _set_run_font(r, size=10.5)

    p = doc.add_paragraph("输出3：自然语言问数结果（Text2SQL，只读查询）：")
    for r in p.runs:
        _set_run_font(r, size=10.5)
    p = doc.add_paragraph(
        "问①\u201c各存放位置库存数量排名\u201d → SQL：SELECT location, COUNT(*) AS n, SUM(stock_qty) AS total_qty "
        "FROM fact_inventory GROUP BY location ORDER BY total_qty DESC LIMIT 5"
    )
    for r in p.runs:
        _set_run_font(r, size=10.5)
    q_rows = [[loc, n, q] for loc, n, q in s["q_loc"]]
    _table(doc, ["存放位置", "物资种类数", "库存总量"], q_rows)
    p = doc.add_paragraph(
        "问②\u201c出库流水有多少条、合计多少数量\u201d → SQL：SELECT flow_type, COUNT(*) AS n, SUM(quantity) AS qty "
        "FROM fact_stock_flow GROUP BY flow_type"
    )
    for r in p.runs:
        _set_run_font(r, size=10.5)
    q2_rows = [[ft, n, q] for ft, n, q in s["flow_by_type"]]
    _table(doc, ["流水类型", "条数", "数量合计"], q2_rows)

    p = doc.add_paragraph(
        "复现方式：python3 scripts/build_demo_env.py 重建干净演示库后可在系统内复现上述全部结果（详见实操演示视频）。"
    )
    for r in p.runs:
        _set_run_font(r, size=10.5)

    doc.save(str(OUT))
    print("已生成:", OUT)
    print("口径: inventory", s["inv_rows"], "行 / flow", s["flow_rows"], "段 / pending", s["flow_pending"])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
